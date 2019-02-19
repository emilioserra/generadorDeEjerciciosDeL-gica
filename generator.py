# -*- coding: utf-8 -*-
"""
Created on 15/02/2019
@author: Emilio Serrano, http://emilioserra.oeg-upm.net/
"""



from random import shuffle
import configparser 
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor


import os


#leer fichero de configuración 
config = configparser.ConfigParser()
config.read('./config.cfg') #usar  .encode('latin1').decode('utf8') para tratar acentos
#load data frame with exercises 
xls_file = pd.ExcelFile('./Ejercicios lógica (respuestas).xlsx') 
df =  xls_file.parse('Respuestas de formulario 1')





#devolver todos los índices de del dataframe
def todosLosIndices():
  return list(range(0, df.shape[0]))





#==============================================================================
# Filtra los ejercicios con criterios como un tema y devuelve índices del excel/dataframe a considerar 
# (con orden aleatorio y máximo de ejercicios fijado en fichero de configuración)
# recibe como parámetro la sección del fichero de configuración, ejemplo [FORMALIZACION]
# Nota 1: se puede añdir filtro extra como considerar sólo ejercicios con solución
#==============================================================================


def indicesFiltradosPorTemaYAleatorios(configSection):
  temas = config[configSection]['TEMAS'] 
  #errores de codificación por acentos
  temas= temas.encode('latin1').decode('utf8') 
  temas =   temas.split(",")
  indexes = []
  
  print("FILTRANDO EJERCICIOS PARA " + configSection + "...")
  
  #leer la lista de temas a incluir en config, recuperar índices de dataframe y acumular en lista
  for t in temas:
    print("Índices en excel de ejercicios tema ", t)
    indexesForOneModule=list((df[df['Tema']==t]).index.values)
    print(indexesForOneModule)
    indexes.extend(indexesForOneModule)


  #reordenar aleatoriamente y quedarse con el máximo considerado en el fichero de config
  indexesOrderedAndFiltered = list(indexes)
  shuffle(indexesOrderedAndFiltered)  
  indexesOrderedAndFiltered = indexesOrderedAndFiltered[0:int(config['DEFAULT']['MAX_EXERCISES'])+1]
  print("Índices seleccionados: ", indexesOrderedAndFiltered)
  print("...OK")
  return indexesOrderedAndFiltered



#==============================================================================
# Método llamado en escribirDataFrame para tratar los caracteres especiales del excel antes de escribirlos en las salidas word
# Recorre el string caracter a caracter comprobando si toma estas acciones
# *TEXTO* para resaltar
# | | para centrar texto
#  
# Nota 1: seguro que no es la forma más elegante de hacerlo, especialmente por usar continue, pero switch en python es complicado)
# Nota 2: la negrita daba fallos en los siguientes caracteres: ['∨', '∧', '∃', '∀', '⊨', '⊢', '⊭']. En una primera solución se omitieron. Luego se hizo que se pusiesen en rojo.
# Nota 3: se pueden mapear los caractéres de símbolos lógicos (¬, ∨, ∧, → ,↔, ∃x, ∀x, ⊨, ⊢) a versiones word más elegantes y con negrita. Igualmente se puede "escupir" código latex en el word
# 
#==============================================================================
    
    
def tratamientoCaracteresEspeciales(document, texto):
    resaltadoAbierto=False
    centradoAbierto=False
    #estos caracteres dan error al poner en negrita en word
    #caracteresSinNegrita=['∨', '∧', '∃', '∀', '⊨', '⊢', '⊭']

    p = document.add_paragraph()
    
    for c in texto:
        
        if c=='*' and resaltadoAbierto==False:
            resaltadoAbierto=True
            continue
        if c=='*' and resaltadoAbierto==True:            
            resaltadoAbierto=False
            continue       
        if c=='|' and centradoAbierto==False:            
            p = document.add_paragraph() #parrafo nuevo centrado
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER    
            centradoAbierto=True
            continue  
    
        if c=='#':
            p.add_run('\t')
            continue  
        
        
        
        
        if c=='|' and centradoAbierto==True:            
            centradoAbierto=False
            p = document.add_paragraph() #parrafo nuevo sin centrar
            continue  
        
        if(resaltadoAbierto): 
            font = p.add_run(c).font
            font.color.rgb = RGBColor(0xFF, 0x00, 0x00) #color resaltado RGB, rojo
         
        else:
            font = p.add_run(c).font
            font.color.rgb = RGBColor(0x00, 0x00, 0x00)#color por defecto, negro
        

#==============================================================================
# Método para escribir los ficheros con enunciados y enunciados con soluciones. Recibe lista de índices de ejercicios a considerar. 
# También genera un doc con los códigos de los ejercicios utilizados. Código 0 significa primer ejercicio del excel (código+2 la fila en el excel)
# Se debe pasar una lista de índices del dataframe para escribir en los ficheros.
# Esta lista puede venir de filtrarUnTema, ver método, o la lista completa si se quiere generar un fichero de pruebas 
# 
#==============================================================================
def escribirDataFrame(configSection, indexes, withSolutions):
    
    #all exercises for the list of indexes
    #print(df.iloc[indexes])
    
    print("GENERANDO FICHERO DE " + configSection + ", ¿soluciones incluidas?: " + str(withSolutions) + "...") 
    document = Document()  
    #añadir cabecera
    document.add_heading( config[configSection]['OUTPUT_TITLE'].encode('latin1').decode('utf8')        , 0)
    if withSolutions:   
        document.add_heading('(Con soluciones)', level=1)

    #bucle para cada ejercicio
    #loop: i goes from 0 to the number of indexes of exercises in the dataframe
    #si con soluciones es verdadero se añade ese campo del exce.
    for i in range(0, len(indexes)):             
        document.add_heading("Ejercicio " + str(i+1) +"." , level=1)
        exerciseDf= df.iloc[[indexes[i]]] #dataframe con sólo el ejercicio seleccionado
        if withSolutions:   
            document.add_heading("Enunciado" , level=2)
        #document.add_paragraph(exerciseDf['Enunciado'].values[0]) 
        tratamientoCaracteresEspeciales(document,exerciseDf['Enunciado'].values[0])
        if withSolutions:   
            document.add_heading("Solución" , level=2)
            #document.add_paragraph(exerciseDf['Solución'].values[0]) 
            tratamientoCaracteresEspeciales(document,exerciseDf['Solución'].values[0])
            
    
    
 
    
    #escribir fichero word de salida (se añade coletilla si es con soluciones)
    outputFile = config[configSection]['OUTPUT'].encode('latin1').decode('utf8')
    if withSolutions:
        #salvar en nombre de fichero añadiendo "ConSoluciones
        outputFile2 = os.path.splitext(outputFile)[0] + "ConSoluciones" + os.path.splitext(outputFile)[1]         
        document.save(outputFile2)
    else:
        document.save(outputFile)

    
    #escribir fichero de salida con índices de ejercicios generados en la última llamada a este método
    outputFile3 =  os.path.splitext(outputFile)[0] + "ÍndicesDeEjercicios.txt"
    with open(outputFile3, 'w') as f:
        f.write(str(indexes))
    
    print("...OK")

                
if __name__ == '__main__':

   
   #generación de ficheros con todos los ejercicios para revisión
   escribirDataFrame('TODOS',todosLosIndices(), False)
   escribirDataFrame('TODOS',todosLosIndices(), True )
   
   
   
   #generación de ficheros de listado específico
   escribirDataFrame('LISTADO',  config['LISTADO']['LIST'].split(",") , False)
   escribirDataFrame('LISTADO',  config['LISTADO']['LIST'].split(",") , True)
   
   #generación de ejercicios para un tema concreto, configurado en sección de config.cfg
   #se generan primero sin soluciones y luego con ellas
   

   #guardar lista con secciones tras todos y listado para iterar ('LP_FORMALIZACION', 'LP_SEMÁNTICA'...)
   seccionesDeFicheroDeConfig= config.sections()
   seccionesDeFicheroDeConfig = seccionesDeFicheroDeConfig[2:len(seccionesDeFicheroDeConfig)]
   for s in seccionesDeFicheroDeConfig:
       indices=indicesFiltradosPorTemaYAleatorios(s)
       escribirDataFrame(s,indices, False)
       escribirDataFrame(s,indices, True )
   
        
   
    
 
    
 
 
   